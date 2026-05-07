#!/usr/bin/env python3
"""Convert CN6000 dissertation plain text to a Word document (.docx)."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def is_dash_line(s: str) -> bool:
    t = s.strip()
    return bool(t) and bool(re.match(r"^[-=]+$", t)) and len(t) >= 3


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    src = root / "CN6000_Dissertation_Draft_AmritTamang_2546484.txt"
    if len(sys.argv) >= 2:
        src = Path(sys.argv[1])
    out = src.with_suffix(".docx")

    lines = src.read_text(encoding="utf-8").splitlines()
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        if is_dash_line(s):
            i += 1
            continue

        if re.match(r"^CHAPTER \d+:", s, re.IGNORECASE):
            doc.add_heading(s, level=1)
            i += 1
            continue

        if re.match(r"^Chapter \d+:", s):
            doc.add_heading(s, level=1)
            i += 1
            continue

        if re.match(r"^Appendix [A-H]\b", s, re.IGNORECASE):
            doc.add_heading(s, level=1)
            i += 1
            continue

        if re.match(r"^REFERENCE LIST", s, re.IGNORECASE):
            doc.add_heading(s, level=1)
            i += 1
            continue

        if s == "END OF DOCUMENT":
            doc.add_paragraph(s)
            i += 1
            continue

        if re.match(r"^\d+\.\d+\.\d+\s+\S", s):
            doc.add_heading(s, level=3)
            i += 1
            if i < n and is_dash_line(lines[i]):
                i += 1
            continue

        if re.match(r"^\d+\.\d+\s+\S", s):
            doc.add_heading(s, level=2)
            i += 1
            if i < n and is_dash_line(lines[i]):
                i += 1
            continue

        if s.isupper() and len(s) < 42 and ".." not in s and "|" not in s:
            doc.add_heading(s, level=1)
            i += 1
            continue

        if s.startswith("TITLE PAGE NOTES"):
            doc.add_heading(s, level=2)
            i += 1
            continue

        block = [line]
        i += 1
        while i < n and lines[i].strip():
            ns = lines[i].strip()
            if is_dash_line(ns):
                break
            if re.match(r"^CHAPTER \d+:", ns, re.IGNORECASE):
                break
            if re.match(r"^Chapter \d+:", ns):
                break
            if re.match(r"^Appendix [A-H]\b", ns, re.IGNORECASE):
                break
            if ns.startswith("REFERENCE LIST") or ns.upper().startswith(
                "REFERENCE LIST"
            ):
                break
            if re.match(r"^\d+\.\d+\.\d+\s+\S", ns):
                break
            if re.match(r"^\d+\.\d+\s+\S", ns):
                break
            if ns.isupper() and len(ns) < 42 and ".." not in ns and "|" not in ns:
                break
            block.append(lines[i])
            i += 1

        text = "\n".join(block).rstrip()
        if text:
            doc.add_paragraph(text)

    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
